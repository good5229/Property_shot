#!/usr/bin/env python3
"""Build evaluator-facing submission DOCX files from Markdown sources.

The game guide uses the documents skill's compact_reference_guide preset.
The AI report uses standard_business_brief. Both use editorial_cover and a
named Korean-font override: embedded NanumGothic replaces Calibri so Hangul
renders consistently in LibreOffice and the final PDF.
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage

import build_final_report_docx as base


ROOT = Path(__file__).resolve().parents[2]
FONT = base.FONT
BLUE = base.BLUE
DARK_BLUE = base.DARK_BLUE
NAVY = base.NAVY
MUTED = base.MUTED
GOLD = base.GOLD
DOCUMENT_LANGUAGE = "ko-KR"


@dataclass(frozen=True)
class ReportConfig:
    source: Path
    output: Path
    title: str
    kicker: str
    running_label: str
    preset: str
    header_fill: str
    tagline: str
    cover_image: Path


def set_document_language(doc: Document) -> None:
    """Declare Korean as the document and inherited text language."""

    doc.core_properties.language = DOCUMENT_LANGUAGE
    settings = doc.settings._element
    theme_lang = settings.find(qn("w:themeFontLang"))
    if theme_lang is None:
        theme_lang = OxmlElement("w:themeFontLang")
        settings.append(theme_lang)
    theme_lang.set(qn("w:val"), DOCUMENT_LANGUAGE)
    theme_lang.set(qn("w:eastAsia"), DOCUMENT_LANGUAGE)

    for style in doc.styles:
        rpr = style._element.get_or_add_rPr()
        lang = rpr.find(qn("w:lang"))
        if lang is None:
            lang = OxmlElement("w:lang")
            rpr.append(lang)
        lang.set(qn("w:val"), DOCUMENT_LANGUAGE)
        lang.set(qn("w:eastAsia"), DOCUMENT_LANGUAGE)


CONFIGS = {
    "game": ReportConfig(
        source=ROOT / "report" / "game_introduction.md",
        output=ROOT / "report" / "dist" / "property_shot_game_guide.docx",
        title="속성 한방(Property Shot)",
        kicker="Game Guide",
        running_label="Property Shot · Game Guide",
        preset="compact_reference_guide",
        header_fill="E8EEF5",
        tagline="속성을 옮기고, 실패까지 다음 해법으로 바꾸는 물리 퍼즐",
        cover_image=ROOT / "screenshots" / "commercial-vertical-slice" / "390x844-current-play-audit.png",
    ),
    "ai": ReportConfig(
        source=ROOT / "report" / "ai_technical_report.md",
        output=ROOT / "report" / "dist" / "property_shot_ai_technical_report.docx",
        title="속성 한방(Property Shot)",
        kicker="AI Technical Report",
        running_label="Property Shot · AI Technical Report",
        preset="standard_business_brief",
        header_fill="F2F4F7",
        tagline="프롬프트를 제작 계약으로 바꾸고, 독립 감사로 검증한 AI 협업",
        cover_image=ROOT / "test" / "goldens" / "difficulty_easy_first_arrival_390x844.png",
    ),
    "portfolio": ReportConfig(
        source=ROOT / "report" / "portfolio.md",
        output=ROOT / "report" / "dist" / "property_shot_portfolio.docx",
        title="속성 한방(Property Shot)",
        kicker="GAME · AI · PRODUCT CASE STUDY",
        running_label="Property Shot · Portfolio",
        preset="compact_reference_guide",
        header_fill="E8F4F1",
        tagline="게임 디렉팅, AI 오케스트레이션, 검증 가능한 제품 제작",
        cover_image=ROOT / "screenshots" / "commercial-vertical-slice" / "stage4-property-ready.png",
    ),
}


def configure_styles(doc: Document, config: ReportConfig) -> dict[str, float]:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    if config.preset == "compact_reference_guide":
        tokens = {
            "body_after": 6,
            "body_line": 1.25,
            "h1_before": 18,
            "h1_after": 10,
            "h2_before": 14,
            "h2_after": 7,
            "h3_before": 10,
            "h3_after": 5,
            "list_left": 0.375,
            "list_hanging": 0.188,
            "list_after": 4,
            "list_line": 1.25,
        }
    else:
        tokens = {
            "body_after": 6,
            "body_line": 1.10,
            "h1_before": 16,
            "h1_after": 8,
            "h2_before": 12,
            "h2_after": 6,
            "h3_before": 8,
            "h3_after": 4,
            "list_left": 0.5,
            "list_hanging": 0.25,
            "list_after": 8,
            "list_line": 1.167,
        }

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal._element.rPr.rFonts.set(qn(f"w:{key}"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(tokens["body_after"])
    normal.paragraph_format.line_spacing = tokens["body_line"]
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    headings = [
        ("Heading 1", 16, BLUE, tokens["h1_before"], tokens["h1_after"]),
        ("Heading 2", 13, BLUE, tokens["h2_before"], tokens["h2_after"]),
        ("Heading 3", 12, DARK_BLUE, tokens["h3_before"], tokens["h3_after"]),
    ]
    for name, size, color, before, after in headings:
        style = styles[name]
        style.font.name = FONT
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            style._element.rPr.rFonts.set(qn(f"w:{key}"), FONT)
        style._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    return tokens


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("Page ")
    base.set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def add_cover(doc: Document, config: ReportConfig) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hr = hp.add_run(config.running_label)
    base.set_run_font(hr, size=9, color=MUTED)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    base.paragraph_bottom_border(hp, "D7DBE2", "6")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run(config.running_label + " · ")
    base.set_run_font(fr, size=8.5, color=MUTED)
    add_page_field(fp)

    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    base.set_run_font(kicker.add_run(config.kicker), size=11, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    base.set_run_font(title.add_run(config.title), size=30, color=NAVY, bold=True)

    strap = doc.add_paragraph()
    strap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    strap.paragraph_format.space_after = Pt(18)
    base.set_run_font(
        strap.add_run(config.tagline),
        size=10.5,
        color=GOLD,
    )

    hero = doc.add_paragraph()
    hero.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hero.paragraph_format.space_after = Pt(18)
    picture = hero.add_run().add_picture(str(config.cover_image), width=Inches(1.7))
    picture._inline.docPr.set("descr", "속성 한방 실제 플레이 화면")
    picture._inline.docPr.set("title", "Property Shot gameplay")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(6)
    base.set_run_font(meta.add_run("2026-08-10 KST"), size=11, color=NAVY, bold=True)

    event = doc.add_paragraph()
    event.alignment = WD_ALIGN_PARAGRAPH.CENTER
    event.paragraph_format.space_after = Pt(10)
    base.set_run_font(event.add_run("NAN 2026 Game × AI 해커톤 사전 과제"), size=10, color=MUTED)
    doc.add_page_break()


def add_workflow_diagram(doc: Document) -> None:
    """Add an editable, left-to-right production workflow."""
    labels = [
        "사용자\n피드백",
        "→",
        "기능 계약",
        "→",
        "전문 실행",
        "→",
        "독립 감사",
        "→",
        "증거·통합",
    ]
    table = doc.add_table(rows=1, cols=len(labels))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, (cell, label) in enumerate(zip(table.rows[0].cells, labels)):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        width = Inches(1.02 if index % 2 == 0 else 0.24)
        cell.width = width
        cell._tc.tcPr.tcW.set(qn("w:w"), str(round(width.inches * 1440)))
        cell._tc.tcPr.tcW.set(qn("w:type"), "dxa")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(8)
        run = paragraph.add_run(label)
        base.set_run_font(
            run,
            size=9.5 if index % 2 == 0 else 15,
            color=NAVY if index % 2 == 0 else GOLD,
            bold=True,
        )
        if index % 2 == 0:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "E8F4F1")
            cell._tc.get_or_add_tcPr().append(shading)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(10)
    base.set_run_font(note.add_run("감사에서 문제가 발견되면 기능 계약 단계로 돌아간다."), size=8.5, color=MUTED)


def next_num_id(doc: Document) -> int:
    root = doc.part.numbering_part.element
    values = [int(node.get(qn("w:numId"))) for node in root.findall(qn("w:num"))]
    return max(values, default=0) + 1


def add_numbering(doc: Document, *, bullet: bool, tokens: dict[str, float]) -> int:
    root = doc.part.numbering_part.element
    num_id = next_num_id(doc)
    abstract_id = num_id + 100

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(round(tokens["list_left"] * 1440)))
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(round(tokens["list_left"] * 1440)))
    ind.set(qn("w:hanging"), str(round(tokens["list_hanging"] * 1440)))
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    root.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), str(abstract_id))
    num.append(ref)
    root.append(num)
    return num_id


def add_list_paragraph(
    doc: Document,
    text: str,
    num_id: int,
    tokens: dict[str, float],
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(tokens["list_after"])
    p.paragraph_format.line_spacing = tokens["list_line"]
    ppr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    ppr.append(num_pr)
    base.add_inline_runs(p, text)


def add_body_paragraph(doc: Document, text: str, tokens: dict[str, float]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(tokens["body_after"])
    p.paragraph_format.line_spacing = tokens["body_line"]
    base.add_inline_runs(p, text)


def add_report_image(doc: Document, alt: str, rel_path: str) -> None:
    """Add an image at a readable size without overflowing portrait pages."""
    img_path = (ROOT / "report" / rel_path).resolve()
    if not img_path.exists():
        p = doc.add_paragraph()
        base.add_inline_runs(p, f"[이미지 누락: {alt} - {rel_path}]")
        return

    with PILImage.open(img_path) as image:
        width_px, height_px = image.size
    portrait = height_px / max(width_px, 1) > 1.35
    width = Inches(2.35 if portrait else 4.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(img_path), width=width)
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", alt)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    base.set_run_font(cap.add_run(alt), size=8.5, color=MUTED)


def build(config: ReportConfig) -> Path:
    markdown = config.source.read_text(encoding="utf-8")
    doc = Document()
    set_document_language(doc)
    tokens = configure_styles(doc, config)
    add_cover(doc, config)
    base.LIGHT_FILL = config.header_fill

    lines = base.parse_body_lines(markdown)
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        if not line.strip():
            idx += 1
            continue
        if line.startswith("```"):
            idx += 1
            code = []
            while idx < len(lines) and not lines[idx].startswith("```"):
                code.append(lines[idx])
                idx += 1
            base.add_code_block(doc, code)
            idx += 1
            continue
        if line.strip() == "[[WORKFLOW]]":
            add_workflow_diagram(doc)
            idx += 1
            continue
        if line.startswith("|") and idx + 1 < len(lines) and lines[idx + 1].startswith("|"):
            rows = [base.split_table_row(line)]
            idx += 2
            while idx < len(lines) and lines[idx].startswith("|"):
                rows.append(base.split_table_row(lines[idx]))
                idx += 1
            base.add_table(doc, rows)
            continue
        image = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if image:
            add_report_image(doc, image.group(1), image.group(2))
            idx += 1
            continue
        heading = re.match(r"^(#{2,4})\s+(.*)$", line)
        if heading:
            base.add_heading(doc, heading.group(2), len(heading.group(1)) - 1)
            idx += 1
            continue
        if re.match(r"^-\s+", line):
            num_id = add_numbering(doc, bullet=True, tokens=tokens)
            while idx < len(lines):
                item = re.match(r"^-\s+(.*)$", lines[idx])
                if not item:
                    break
                add_list_paragraph(doc, item.group(1), num_id, tokens)
                idx += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            num_id = add_numbering(doc, bullet=False, tokens=tokens)
            while idx < len(lines):
                item = re.match(r"^\d+\.\s+(.*)$", lines[idx])
                if not item:
                    break
                add_list_paragraph(doc, item.group(1), num_id, tokens)
                idx += 1
            continue
        if line.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(tokens["body_after"])
            base.paragraph_bottom_border(p, "D7DBE2", "4")
            base.add_inline_runs(p, line.lstrip("> ").strip(), color=MUTED)
            idx += 1
            continue
        add_body_paragraph(doc, line, tokens)
        idx += 1

    config.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(config.output)
    base.embed_nanum_fonts(config.output)
    return config.output


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--kind", choices=sorted(CONFIGS), required=True)
    args = parser.parse_args()
    print(build(CONFIGS[args.kind]))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
