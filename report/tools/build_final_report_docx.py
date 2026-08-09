#!/usr/bin/env python3
"""Build the final Property Shot report DOCX from report/final_project_report.md.

The document uses the documents skill's narrative_proposal preset with one
Korean-font override: NanumGothic is applied where Calibri would not reliably
render Hangul on the local macOS/LibreOffice renderer.
"""

from __future__ import annotations

import re
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Iterable
from uuid import UUID, uuid4
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor
from docx.text.run import Run


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "report" / "final_project_report.md"
OUTPUT = ROOT / "report" / "dist" / "property_shot_final_report.docx"
REGULAR_FONT_FILE = ROOT / "assets" / "fonts" / "NanumGothic-Regular.ttf"
BOLD_FONT_FILE = ROOT / "assets" / "fonts" / "NanumGothic-Bold.ttf"

FONT = "NanumGothic"
MONO_FONT = "Menlo"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(88, 98, 112)
GOLD = RGBColor(122, 90, 0)
LIGHT_FILL = "F4F6F9"
BORDER = "D7DBE2"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
FONT_REL_TYPE = f"{R_NS}/font"
OBFUSCATED_FONT_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.obfuscatedFont"
)

ET.register_namespace("w", W_NS)
ET.register_namespace("r", R_NS)
ET.register_namespace("", PKG_REL_NS)


def set_run_font(
    run,
    *,
    name: str = FONT,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def obfuscate_embedded_font(font_path: Path, font_key: UUID) -> bytes:
    """Return an ECMA-376 obfuscated font payload for a DOCX package."""

    payload = bytearray(font_path.read_bytes())
    key = font_key.bytes[::-1]
    for index in range(min(32, len(payload))):
        payload[index] ^= key[index % 16]
    return bytes(payload)


def embed_nanum_fonts(docx_path: Path) -> None:
    """Embed the Korean regular/bold faces so rendering is installation-free."""

    regular_key = uuid4()
    bold_key = uuid4()
    replacements: dict[str, bytes] = {}

    with ZipFile(docx_path, "r") as archive:
        # Preserve python-docx's original namespace declarations byte-for-byte;
        # LibreOffice is strict about the mc:Ignorable prefix map in this part.
        font_table = archive.read("word/fontTable.xml").decode("UTF-8")
        font_entry = f'''  <w:font w:name="{FONT}">
    <w:family w:val="swiss"/>
    <w:pitch w:val="variable"/>
    <w:embedRegular r:id="rId1" w:fontKey="{{{str(regular_key).upper()}}}"/>
    <w:embedBold r:id="rId2" w:fontKey="{{{str(bold_key).upper()}}}"/>
  </w:font>
'''
        replacements["word/fontTable.xml"] = font_table.replace(
            "</w:fonts>", font_entry + "</w:fonts>"
        ).encode("UTF-8")

        rels_name = "word/_rels/fontTable.xml.rels"
        replacements[rels_name] = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="{PKG_REL_NS}">
  <Relationship Id="rId1" Type="{FONT_REL_TYPE}" Target="fonts/font1.odttf"/>
  <Relationship Id="rId2" Type="{FONT_REL_TYPE}" Target="fonts/font2.odttf"/>
</Relationships>'''.encode("UTF-8")

        content_types_name = "[Content_Types].xml"
        content_types = archive.read(content_types_name).decode("UTF-8")
        if 'Extension="odttf"' not in content_types:
            entry = (
                f'<Default Extension="odttf" '
                f'ContentType="{OBFUSCATED_FONT_CONTENT_TYPE}"/>'
            )
            content_types = content_types.replace("</Types>", entry + "</Types>")
        replacements[content_types_name] = content_types.encode("UTF-8")

        replacements["word/fonts/font1.odttf"] = (
            obfuscate_embedded_font(REGULAR_FONT_FILE, regular_key)
        )
        replacements["word/fonts/font2.odttf"] = (
            obfuscate_embedded_font(BOLD_FONT_FILE, bold_key)
        )

        temp_path = docx_path.with_name(f".{docx_path.name}.embedding")
        with ZipFile(temp_path, "w", ZIP_DEFLATED) as output:
            replaced_names = set(replacements)
            for info in archive.infolist():
                if info.filename not in replaced_names:
                    output.writestr(info, archive.read(info.filename))
            for name, payload in replacements.items():
                output.writestr(name, payload)

    temp_path.replace(docx_path)


def set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    add_inline_runs(p, clean_markdown(text), bold_default=bold, color=color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def mark_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def paragraph_bottom_border(paragraph, color: str = BORDER, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_hyperlink(paragraph, label: str, url: str, *, bold: bool = False) -> None:
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run_element = OxmlElement("w:r")
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)
    run = Run(run_element, paragraph)
    set_run_font(run, color=BLUE, bold=bold)
    run.font.underline = True
    run.text = label


def add_inline_runs(paragraph, text: str, *, bold_default: bool = False, color: RGBColor | None = None) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, color=color, bold=bold_default)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name=MONO_FONT, color=color, size=9.5)
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url, bold=bold_default)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, color=color, bold=bold_default)


def clean_markdown(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def table_widths(headers: list[str]) -> list[int]:
    if headers == ["구성", "버전/제약", "용도", "라이선스·공식 출처"]:
        return [1600, 1600, 2600, 3560]
    if headers == ["자산", "출처", "라이선스", "제품 bundle"]:
        return [1700, 2500, 2260, 2900]
    count = len(headers)
    if count == 2:
        return [2700, 6660]
    if count == 3:
        return [1700, 3400, 4260]
    if count == 4:
        return [1200, 2460, 2860, 2840]
    return [9360 // count for _ in range(count)]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    widths = table_widths(rows[0])
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    mark_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        prevent_row_split(table.rows[r_idx])
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
                set_cell_text(cell, value, bold=True, color=NAVY)
            else:
                set_cell_text(cell, value)
def add_cover(doc: Document, lines: list[str]) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = ""
    run = header_p.add_run("Property Shot Final Report")
    set_run_font(run, size=9, color=MUTED)
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Conditional Go · 자동 검증 완료 · 실기기/외부 사용자/법무 증거 미완료")
    set_run_font(run, size=8.5, color=MUTED)

    for _ in range(5):
        doc.add_paragraph()

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(12)
    run = kicker.add_run("Project Result Report")
    set_run_font(run, size=11, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("속성 한방(Property Shot)")
    set_run_font(run, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    run = subtitle.add_run("속성을 옮기고, 비워진 원본과 잔류 공까지 해법으로 쓰는 세로형 물리 퍼즐")
    set_run_font(run, size=13.5, color=DARK_BLUE)

    rule = doc.add_paragraph()
    paragraph_bottom_border(rule, "D7DBE2", "10")
    rule.paragraph_format.space_after = Pt(14)

    meta = [
        ("작성일", "2026-08-09 KST"),
        ("최종 점검", "2026-08-09 11:56:08 KST"),
        ("기준", "main / 30d018a72745a92ccb75276d6ae662303559e3ef"),
        ("판정", "Conditional Go"),
        ("Web 데모", "[바로 플레이](https://good5229.github.io/Property_shot/)"),
        ("저장소", "[GitHub 저장소](https://github.com/good5229/Property_shot)"),
        (
            "배포 기록",
            "[GitHub Actions 성공 기록](https://github.com/good5229/Property_shot/actions/runs/31286806317)",
        ),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    set_table_geometry(table, [1900, 7460])
    mark_table_header(table.rows[0])
    for row_idx, (label, value) in enumerate(meta):
        left, right = table.rows[row_idx].cells
        set_cell_shading(left, LIGHT_FILL)
        set_cell_text(left, label, bold=True, color=NAVY)
        set_cell_text(right, value)

    pic = ROOT / "report" / "assets" / "captures" / "latest-pages" / "390x844-home.png"
    if pic.exists():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(22)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shape = p.add_run().add_picture(str(pic), width=Inches(2.25))
        shape._inline.docPr.set("descr", "최신 공개 데모 홈 화면")
        shape._inline.docPr.set("title", "최신 공개 데모 홈 화면")
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(0)
        run = cap.add_run("최신 공개 데모 홈 화면")
        set_run_font(run, size=8.5, color=MUTED)

    doc.add_page_break()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def parse_body_lines(markdown: str) -> list[str]:
    lines = markdown.splitlines()
    for idx, line in enumerate(lines):
        if line.startswith("## "):
            return lines[idx:]
    return lines


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    add_inline_runs(p, text)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.333
    add_inline_runs(p, text)


def add_list_item(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    add_inline_runs(p, text)


def add_numbered_item(doc: Document, number: str, text: str) -> None:
    # Literal numbering preserves each independent Markdown list's restart.
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    prefix = p.add_run(f"{number}. ")
    set_run_font(prefix)
    add_inline_runs(p, text)


def add_code_block(doc: Document, code: Iterable[str]) -> None:
    for line in code:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        set_run_font(run, name=MONO_FONT, size=9, color=RGBColor(40, 45, 52))
    doc.paragraphs[-1].paragraph_format.space_after = Pt(8)


def add_image(doc: Document, alt: str, rel_path: str) -> None:
    img_path = (ROOT / "report" / rel_path).resolve()
    if not img_path.exists():
        add_paragraph(doc, f"[이미지 누락: {alt} - {rel_path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    width = Inches(2.35) if "390x844" in img_path.name else Inches(4.5)
    shape = p.add_run().add_picture(str(img_path), width=width)
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", alt)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(alt)
    set_run_font(run, size=8.5, color=MUTED)


def build_docx() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    add_cover(doc, markdown.splitlines())

    section = doc.sections[0]
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ""
    run = footer_p.add_run("Property Shot Final Report")
    set_run_font(run, size=8.5, color=MUTED)
    footer_p.add_run(" · ")
    page_run = footer_p.add_run("Page ")
    set_run_font(page_run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    page_run._r.append(fld_char1)
    page_run._r.append(instr_text)
    page_run._r.append(fld_char2)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    lines = parse_body_lines(markdown)
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        if not line.strip():
            idx += 1
            continue
        if line.startswith("```"):
            idx += 1
            code = []
            while idx < len(lines) and not lines[idx].startswith("```"):
                code.append(lines[idx])
                idx += 1
            add_code_block(doc, code)
            idx += 1
            continue
        if line.startswith("|") and idx + 1 < len(lines) and lines[idx + 1].startswith("|"):
            rows = []
            rows.append(split_table_row(line))
            idx += 2
            while idx < len(lines) and lines[idx].startswith("|"):
                rows.append(split_table_row(lines[idx]))
                idx += 1
            add_table(doc, rows)
            continue
        image = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if image:
            add_image(doc, image.group(1), image.group(2))
            idx += 1
            continue
        heading = re.match(r"^(#{2,4})\s+(.*)$", line)
        if heading:
            add_heading(doc, heading.group(2), len(heading.group(1)) - 1)
            idx += 1
            continue
        bullet = re.match(r"^-\s+(.*)$", line)
        if bullet:
            add_list_item(doc, bullet.group(1))
            idx += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.*)$", line)
        if numbered:
            add_numbered_item(doc, numbered.group(1), numbered.group(2))
            idx += 1
            continue
        if line.startswith(">"):
            add_paragraph(doc, line.lstrip("> ").strip())
            idx += 1
            continue
        add_paragraph(doc, line)
        idx += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    embed_nanum_fonts(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
